import json
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_toc_field(paragraph):
    r"""
    Inserts a Table of Contents field into the paragraph.
    Field Code: { TOC \o "1-5" \h \z \u }
    """
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    run = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-5" \\h \\z \\u '
    run._r.append(instrText)

    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar)

    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

def add_toa_field(paragraph, category=1):
    r"""
    Inserts a Table of Authorities field into the paragraph.
    Field Code: { TOA \h \c "1" \p }
    """
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    run = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' TOA \\h \\c "{category}" \\p '
    run._r.append(instrText)

    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar)

    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

def apply_strict_style(paragraph, run, style_def):
    if not style_def:
        return

    # Font Settings
    if "FONT" in style_def:
        run.font.name = style_def["FONT"]
    if "FONT_SIZE" in style_def:
        run.font.size = Pt(style_def["FONT_SIZE"])
    if "BOLD" in style_def:
        run.bold = style_def["BOLD"]
    if "ITALIC" in style_def:
        run.italic = style_def["ITALIC"]
    if "UNDERLINE" in style_def:
        run.underline = style_def["UNDERLINE"]
    if "CAPS" in style_def:
        if style_def["CAPS"] == "ALL":
            run.font.all_caps = True
        elif style_def["CAPS"] == "SMALL":
            run.font.small_caps = True

    # Paragraph Settings
    align_map = {
        "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
        "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
        "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
        "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    if "ALIGNMENT" in style_def:
        paragraph.alignment = align_map.get(style_def["ALIGNMENT"], WD_ALIGN_PARAGRAPH.LEFT)
    
    if "SPACING_AFTER" in style_def:
        paragraph.paragraph_format.space_after = Pt(style_def["SPACING_AFTER"])
    if "SPACING_BEFORE" in style_def:
        paragraph.paragraph_format.space_before = Pt(style_def["SPACING_BEFORE"])
    if "LINE_SPACING" in style_def:
        paragraph.paragraph_format.line_spacing = style_def["LINE_SPACING"]
    
    # Indentation
    if "INDENT_LEFT" in style_def:
        paragraph.paragraph_format.left_indent = Inches(style_def["INDENT_LEFT"])
    if "INDENT_RIGHT" in style_def:
        paragraph.paragraph_format.right_indent = Inches(style_def["INDENT_RIGHT"])
    if "FIRST_LINE_INDENT" in style_def:
        paragraph.paragraph_format.first_line_indent = Inches(style_def["FIRST_LINE_INDENT"])

    # Outline Level (Crucial for TOC)
    # Maps HEADING_1 -> Level 1 (0), HEADING_2 -> Level 2 (1), etc.
    # This allows TOC to pick up the heading even if the style name isn't standard Word "Heading 1"
    if "OUTLINE_LEVEL" in style_def:
        paragraph.paragraph_format.outline_level = style_def["OUTLINE_LEVEL"]
    
    # Pagination (Keep with next, etc.)
    if "PAGINATION" in style_def:
        pg = style_def["PAGINATION"]
        if "KEEP_WITH_NEXT" in pg:
            paragraph.paragraph_format.keep_with_next = pg["KEEP_WITH_NEXT"]
        if "KEEP_LINES_TOGETHER" in pg:
            paragraph.paragraph_format.keep_together = pg["KEEP_LINES_TOGETHER"]
        if "PAGE_BREAK_BEFORE" in pg:
            paragraph.paragraph_format.page_break_before = pg["PAGE_BREAK_BEFORE"]
        if "WIDOW_CONTROL" in pg:
            paragraph.paragraph_format.widow_control = pg["WIDOW_CONTROL"]

    # Tabs
    if "TABS" in style_def:
        from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
        tab_stops = paragraph.paragraph_format.tab_stops
        # Clear existing tabs if any (optional, but safer for strict control)
        # for tab in tab_stops: del tab_stops[0] 
        
        for tab in style_def["TABS"]:
            pos = Inches(tab["POSITION"])
            
            align_map = {
                "LEFT": WD_TAB_ALIGNMENT.LEFT,
                "CENTER": WD_TAB_ALIGNMENT.CENTER,
                "RIGHT": WD_TAB_ALIGNMENT.RIGHT,
                "DECIMAL": WD_TAB_ALIGNMENT.DECIMAL,
                "BAR": WD_TAB_ALIGNMENT.BAR
            }
            alignment = align_map.get(tab.get("ALIGNMENT", "LEFT"), WD_TAB_ALIGNMENT.LEFT)
            
            leader_map = {
                "SPACES": WD_TAB_LEADER.SPACES,
                "DOTS": WD_TAB_LEADER.DOTS,
                "DASHES": WD_TAB_LEADER.DASHES,
                "LINES": WD_TAB_LEADER.LINES
            }
            leader = leader_map.get(tab.get("LEADER", "SPACES"), WD_TAB_LEADER.SPACES)
            
            tab_stops.add_tab_stop(pos, alignment, leader)

def build_document(config_path, styles_path):
    with open(config_path, 'r') as f:
        config = json.load(f)
    
    with open(styles_path, 'r') as f:
        strict_styles = json.load(f)

    doc = Document()
    
    # Apply Document Settings (Margins)
    doc_settings = strict_styles.get("DOCUMENT_SETTINGS", {})
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(doc_settings.get("MARGIN_TOP", 1.0))
        section.bottom_margin = Inches(doc_settings.get("MARGIN_BOTTOM", 1.0))
        section.left_margin = Inches(doc_settings.get("MARGIN_LEFT", 1.0))
        section.right_margin = Inches(doc_settings.get("MARGIN_RIGHT", 1.0))

    styles_map = strict_styles.get("STYLES", {})
    layout = config.get("layout", [])
    
    for block in layout:
        if block["type"] == "paragraph":
            p = doc.add_paragraph()
            text = block.get("text", "")
            run = p.add_run(text)
            
            # Map user-friendly style names to strict style keys if needed
            # For now, we assume the config uses the strict keys (HEADING_1, NORMAL, etc.)
            # or we map them.
            style_key = block.get("style", "NORMAL").upper().replace(" ", "_")
            
            # Fallback mapping for legacy config names
            legacy_map = {
                "HEADING_1": "HEADING_1",
                "BODY_TEXT": "NORMAL",
                "STANDARD_BODY": "NORMAL",
                "SIGNATURE": "NORMAL", # Or define a signature style
                "SIGNATURE_BLOCK": "NORMAL"
            }
            
            mapped_key = legacy_map.get(style_key, style_key)
            style_def = styles_map.get(mapped_key, styles_map.get("NORMAL"))
            
            apply_strict_style(p, run, style_def)
            
        elif block["type"] == "spacer":
            doc.add_paragraph()

        elif block["type"] == "toc":
            p = doc.add_paragraph()
            # Apply TOC Header style if defined, or just Normal
            if block.get("text"):
                p.add_run(block["text"]).bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph() # Spacer
                p = doc.add_paragraph()
            
            add_toc_field(p)
            
        elif block["type"] == "toa":
            p = doc.add_paragraph()
            if block.get("text"):
                p.add_run(block["text"]).bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph() # Spacer
                p = doc.add_paragraph()
            
            add_toa_field(p, category=block.get("category", 1))

    # Output Path
    script_dir = Path(__file__).parent
    # Output to templates/ directory in the skill root
    output_path = script_dir.parent / 'templates/declaration_template.docx'
    
    # Ensure directory exists
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    doc.save(str(output_path))
    print(f"Created strict template at {output_path}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("config_file", help="Path to JSON config")
    parser.add_argument("--styles", default="skills/legal_styles_strict.json", help="Path to styles JSON")
    args = parser.parse_args()
    build_document(args.config_file, args.styles)
