"""Render a DOCX template by replacing {{PLACEHOLDER}} tokens with JSON data.

- Preserves formatting/footnotes because it only replaces text in body paragraphs and tables.
- Keep tokens on a single line/run when possible to avoid split-run edge cases.
- Requires python-docx: pip install python-docx
"""

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any, Dict, List, Mapping, Optional

try:
    from docx import Document
except ImportError:  # pragma: no cover - dependency check
    print("python-docx not installed. Install with: python -m pip install python-docx")
    sys.exit(1)


TOKEN_PATTERN = re.compile(r"{{([^{}]+)}}")


def _get_value(data: Mapping[str, Any], key: str) -> Any:
    """Resolve dotted paths like PARTIES.APPELLEES from nested data."""
    parts = key.split(".")
    current: Any = data
    for part in parts:
        if isinstance(current, Mapping) and part in current:
            current = current[part]
        else:
            return None
    return current


def _render_value(val: Any) -> str:
    """Convert common structures to strings suitable for direct insertion."""
    if val is None:
        return ""
    if isinstance(val, str):
        return val
    if isinstance(val, (int, float)):
        return str(val)
    if isinstance(val, list):
        # Join lists of dicts/strings with blank lines to preserve blocks
        rendered_items: List[str] = []
        for item in val:
            if isinstance(item, Mapping):
                heading = item.get("HEADING") or item.get("heading")
                text = item.get("TEXT") or item.get("text")
                if heading and text:
                    rendered_items.append(f"{heading}\n{text}")
                else:
                    rendered_items.append(_render_value(item))
            else:
                rendered_items.append(_render_value(item))
        return "\n\n".join(rendered_items)
    if isinstance(val, Mapping):
        # Fallback: JSON dump for nested dicts
        return json.dumps(val, indent=2)
    return str(val)


def build_token_map(data: Mapping[str, Any], mapping: Optional[Mapping[str, str]]) -> Dict[str, str]:
    token_map: Dict[str, str] = {}
    for key in data.keys():
        token_map[key] = _render_value(data[key])
    if mapping:
        for token, data_key in mapping.items():
            token_map[token] = _render_value(_get_value(data, data_key) if isinstance(data_key, str) else None)
    return token_map


def replace_in_runs(text: str, token_map: Mapping[str, str]) -> str:
    new_text = text
    for token, replacement in token_map.items():
        placeholder = f"{{{{{token}}}}}"
        if placeholder in new_text:
            new_text = new_text.replace(placeholder, replacement)
    return new_text


def replace_in_paragraph(paragraph, token_map: Mapping[str, str]):
    # Reassemble run text to avoid split-token issues
    full_text = "".join(run.text for run in paragraph.runs)
    new_text = replace_in_runs(full_text, token_map)
    if new_text != full_text:
        # Clear and replace runs to preserve simple styling on the first run
        while paragraph.runs:
            paragraph.runs[0].text = ""
            paragraph.runs[0].clear()
            paragraph._p.remove(paragraph.runs[0]._r)
        paragraph.add_run(new_text)


def replace_in_table(table, token_map: Mapping[str, str]):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, token_map)


def validate_no_placeholders(doc, strict: bool):
    if not strict:
        return
    remaining = []
    for paragraph in doc.paragraphs:
        for match in TOKEN_PATTERN.findall(paragraph.text):
            remaining.append(match)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for match in TOKEN_PATTERN.findall(paragraph.text):
                        remaining.append(match)
    if remaining:
        raise ValueError(f"Unresolved placeholders: {sorted(set(remaining))}")


def main():
    parser = argparse.ArgumentParser(description="Render DOCX from JSON data and placeholders.")
    parser.add_argument("--template", required=True, help="Path to DOCX template with {{PLACEHOLDER}} tokens.")
    parser.add_argument("--data", required=True, help="Path to JSON data file.")
    parser.add_argument("--output", required=True, help="Path for rendered DOCX output.")
    parser.add_argument("--mapping", help="Optional JSON mapping: template_token -> data_key (dotted paths allowed).")
    parser.add_argument("--strict", action="store_true", help="Fail if any placeholders remain.")
    args = parser.parse_args()

    template_path = Path(args.template)
    data_path = Path(args.data)
    output_path = Path(args.output)
    mapping_path = Path(args.mapping) if args.mapping else None

    if not template_path.exists():
        raise SystemExit(f"Template not found: {template_path}")
    if not data_path.exists():
        raise SystemExit(f"Data file not found: {data_path}")

    with data_path.open("r", encoding="utf-8") as f:
        data = json.load(f)

    mapping: Optional[Mapping[str, str]] = None
    if mapping_path:
        with mapping_path.open("r", encoding="utf-8") as f:
            mapping = json.load(f)

    token_map = build_token_map(data, mapping)

    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, token_map)

    for table in doc.tables:
        replace_in_table(table, token_map)

    validate_no_placeholders(doc, args.strict)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    main()
