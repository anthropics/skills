#!/usr/bin/env python3
"""
Extract text from a .docx manuscript with paragraph indices.

Usage:
    python extract_docx.py <input.docx> [output.txt]

Output format:
    [0] First paragraph text
    [1] Second paragraph text
    ...

Requires: pip install python-docx
"""

import sys
import os

def extract(docx_path, output_path=None):
    try:
        from docx import Document
    except ImportError:
        print("Error: python-docx not installed. Run: pip install python-docx")
        sys.exit(1)

    if not os.path.exists(docx_path):
        print(f"Error: File not found: {docx_path}")
        sys.exit(1)

    if output_path is None:
        base = os.path.splitext(os.path.basename(docx_path))[0]
        output_path = os.path.join(os.path.dirname(docx_path), f"_{base}_text.txt")

    doc = Document(docx_path)
    with open(output_path, "w", encoding="utf-8") as f:
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip():
                f.write(f"[{i}] {p.text}\n\n")

    count = sum(1 for p in doc.paragraphs if p.text.strip())
    print(f"Done. {count} paragraphs extracted to {output_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    extract(sys.argv[1], sys.argv[2] if len(sys.argv) > 2 else None)
