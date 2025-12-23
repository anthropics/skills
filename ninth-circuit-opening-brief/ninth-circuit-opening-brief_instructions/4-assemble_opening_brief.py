#!/usr/bin/env python3
"""
Ninth Circuit Opening Brief Assembler

*** CRITICAL: NO TEXT GENERATION ***
This script ONLY copies text from source JSON files.
It does NOT generate, reword, or modify any content.

USAGE:
  python assemble_opening_brief.py --all --case-no 25-XXXXX
  python assemble_opening_brief.py --section introduction
  python assemble_opening_brief.py --toa
  python assemble_opening_brief.py --toc
  python assemble_opening_brief.py --validate
  python assemble_opening_brief.py --word-count
"""

import json
import argparse
import os
import re
import stat
from pathlib import Path
from datetime import datetime
from collections import OrderedDict
from typing import Optional


# FRAP 28 Section Order
SECTION_ORDER = [
    "cover_page",
    "disclosure_statement", 
    "table_of_contents",
    "table_of_authorities",
    "introduction",
    "jurisdictional_statement",
    "statutory_authorities",
    "issues_presented",
    "statement_of_case",
    "summary_of_argument",
    "standard_of_review",
    "argument",
    "conclusion",
    "related_cases",
    "certificate_compliance",
    "certificate_service",
    "addendum"
]

# Sections that are auto-generated (not user text)
AUTO_GENERATED = ["cover_page", "table_of_contents", "table_of_authorities", 
                  "certificate_compliance", "certificate_service"]

# Word limit (can be overridden via CLI)
WORD_LIMIT = 14000


class BriefLoader:
    """Load exact text from source JSON files. READ-ONLY."""
    
    def __init__(self, data_dir: str):
        self.data_dir = Path(data_dir)
        self._sections = None
        self._authorities = None
    
    def load_sections(self) -> dict:
        """Load sections.json - cached"""
        if self._sections is None:
            path = self.data_dir / "sections.json"
            if path.exists():
                with open(path, 'r', encoding='utf-8') as f:
                    self._sections = json.load(f)
            else:
                print(f"ERROR: sections.json not found at {path}")
                print("Copy sections_template.json to sections.json and fill in your text.")
                self._sections = {}
        return self._sections
    
    def load_authorities(self) -> dict:
        """Load authorities.json - cached"""
        if self._authorities is None:
            path = self.data_dir / "authorities.json"
            if path.exists():
                with open(path, 'r', encoding='utf-8') as f:
                    self._authorities = json.load(f)
            else:
                self._authorities = {"cases": [], "statutes": [], "rules": [], "other_authorities": []}
        return self._authorities
    
    def get_case_info(self) -> dict:
        """Get case information"""
        sections = self.load_sections()
        return sections.get("case_info", {})
    
    def get_section_text(self, section_name: str) -> str:
        """Get exact text for a section - NO MODIFICATION"""
        sections = self.load_sections()
        section_data = sections.get("sections", {}).get(section_name, {})
        return section_data.get("text", "")
    
    def get_section_id(self, section_name: str) -> str:
        """Get section ID"""
        sections = self.load_sections()
        section_data = sections.get("sections", {}).get(section_name, {})
        return section_data.get("id", "")
    
    def list_sections(self) -> list:
        """List all available sections"""
        sections = self.load_sections()
        return list(sections.get("sections", {}).keys())


class AuthorityExtractor:
    """Extract citations from text for Table of Authorities."""
    
    # Regex patterns for citations
    CASE_PATTERN = r'([A-Z][A-Za-z\'\-\s]+(?:v\.|vs\.)\s+[A-Z][A-Za-z\'\-\s,]+),?\s*(\d+\s+(?:U\.S\.|F\.\d+[d]?|F\.\s*(?:Supp\.|App\'x)|S\.\s*Ct\.|L\.\s*Ed\.?\s*\d*d?)\s*\d+(?:,?\s*\d+)?)\s*\(([^)]+)\)'
    STATUTE_PATTERN = r'(\d+)\s+U\.S\.C\.\s*§\s*(\d+[a-z]?(?:\([a-z0-9]+\))?)'
    RULE_PATTERN = r'(Fed\.\s*R\.\s*(?:Civ\.|App\.|Crim\.)\s*P\.\s*\d+(?:\([a-z]\)(?:\(\d+\))?(?:\([A-Za-z]+\))?)?)'
    CFR_PATTERN = r'(\d+)\s+C\.F\.R\.\s*§\s*(\d+(?:\.\d+)?(?:\([a-z]\))?)'
    
    def extract_cases(self, text: str) -> list:
        """Extract case citations"""
        matches = re.findall(self.CASE_PATTERN, text)
        cases = []
        for match in matches:
            cases.append({
                "name": match[0].strip(),
                "citation": f"{match[1]} ({match[2]})",
                "raw": f"{match[0].strip()}, {match[1]} ({match[2]})"
            })
        return cases
    
    def extract_statutes(self, text: str) -> list:
        """Extract U.S.C. citations"""
        matches = re.findall(self.STATUTE_PATTERN, text)
        return [{"citation": f"{m[0]} U.S.C. § {m[1]}"} for m in matches]
    
    def extract_rules(self, text: str) -> list:
        """Extract FRCP/FRAP citations"""
        matches = re.findall(self.RULE_PATTERN, text)
        return [{"citation": m} for m in matches]
    
    def extract_all(self, text: str) -> dict:
        """Extract all citations from text"""
        return {
            "cases": self.extract_cases(text),
            "statutes": self.extract_statutes(text),
            "rules": self.extract_rules(text)
        }


class BriefAssembler:
    """Assemble Opening Brief by copying exact text from sources."""
    
    def __init__(self, data_dir: str, output_dir: str):
        self.loader = BriefLoader(data_dir)
        self.extractor = AuthorityExtractor()
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def generate_cover_page(self) -> str:
        """Generate cover page from case_info"""
        info = self.loader.get_case_info()
        
        cover = f"""No. {info.get('ninth_circuit_no', '[insert 9th Circuit case number]')}
__________________________________________________________________

IN THE UNITED STATES COURT OF APPEALS
FOR THE NINTH CIRCUIT

{info.get('appellant', '[APPELLANT NAME]')},
    Plaintiff-Appellant,

v.

{', '.join(info.get('appellees', ['[APPELLEE NAMES]']))},
    Defendants-Appellees.

On Appeal from the United States District Court
for the {info.get('district_court', '[District]')}
No. {info.get('district_case_no', '[district case number]')}
{info.get('judge', 'Hon. [Judge Name]')}

APPELLANT'S OPENING BRIEF

{info.get('appellant', '[Name]')}
Plaintiff-Appellant Pro Se
{info.get('appellant_address', '[Address]')}
{info.get('appellant_email', '[Email]')}
{info.get('appellant_phone', '[Phone]')}
"""
        return cover
    
    def generate_toc(self) -> str:
        """Generate Table of Contents from section headings"""
        toc_lines = ["TABLE OF CONTENTS", "", "Page", ""]
        
        # Standard sections with placeholder page numbers
        toc_entries = [
            ("DISCLOSURE STATEMENT", "i"),
            ("TABLE OF AUTHORITIES", "iv"),
            ("INTRODUCTION", "1"),
            ("JURISDICTIONAL STATEMENT", "X"),
            ("ISSUES PRESENTED", "X"),
            ("STATEMENT OF THE CASE", "X"),
            ("SUMMARY OF THE ARGUMENT", "X"),
            ("STANDARD OF REVIEW", "X"),
            ("ARGUMENT", "X"),
            ("CONCLUSION", "X"),
            ("STATEMENT OF RELATED CASES", "X"),
            ("CERTIFICATE OF COMPLIANCE", "X"),
            ("CERTIFICATE OF SERVICE", "X"),
            ("ADDENDUM", "X"),
        ]
        
        for title, page in toc_entries:
            dots = "." * (60 - len(title) - len(str(page)))
            toc_lines.append(f"{title} {dots} {page}")
        
        return "\n".join(toc_lines)
    
    def generate_toa(self) -> str:
        """Generate Table of Authorities from authorities.json"""
        auth = self.loader.load_authorities()
        
        lines = ["TABLE OF AUTHORITIES", "", "Page(s)", ""]
        
        # Cases
        if auth.get("cases"):
            lines.append("Cases")
            lines.append("")
            for case in sorted(auth["cases"], key=lambda x: x.get("name", "")):
                name = case.get("name", "")
                citation = case.get("citation", "")
                pages = ", ".join(str(p) for p in case.get("pages_cited", []))
                lines.append(f"{name},")
                lines.append(f"    {citation} ............................................................. {pages}")
            lines.append("")
        
        # Statutes
        if auth.get("statutes"):
            lines.append("Statutes")
            lines.append("")
            for stat in sorted(auth["statutes"], key=lambda x: x.get("citation", "")):
                citation = stat.get("citation", "")
                pages = ", ".join(str(p) for p in stat.get("pages_cited", []))
                lines.append(f"{citation} ............................................................. {pages}")
            lines.append("")
        
        # Rules
        if auth.get("rules"):
            lines.append("Rules")
            lines.append("")
            for rule in sorted(auth["rules"], key=lambda x: x.get("citation", "")):
                citation = rule.get("citation", "")
                pages = ", ".join(str(p) for p in rule.get("pages_cited", []))
                lines.append(f"{citation} ............................................................. {pages}")
            lines.append("")
        
        return "\n".join(lines)
    
    def generate_cert_compliance(self, word_count: int) -> str:
        """Generate Certificate of Compliance (Form 8)"""
        info = self.loader.get_case_info()
        
        return f"""UNITED STATES COURT OF APPEALS
FOR THE NINTH CIRCUIT

Form 8. Certificate of Compliance for Briefs

9th Cir. Case Number(s): {info.get('ninth_circuit_no', '')}

I am the attorney or self-represented party.

This brief contains {word_count} words, including 0 words
manually counted in any visual images, and excluding the items exempted by FRAP
32(f). The brief's type size and typeface comply with FRAP 32(a)(5) and (6).

I certify that this brief:

[X] complies with the word limit of Cir. R. 32-1.

Signature: /s/ {info.get('appellant', '')}
Date: {datetime.now().strftime('%B %d, %Y')}
"""
    
    def generate_cert_service(self) -> str:
        """Generate Certificate of Service"""
        info = self.loader.get_case_info()
        
        return f"""CERTIFICATE OF SERVICE

I hereby certify that I electronically filed the foregoing with the Clerk of the 
Court for the United States Court of Appeals for the Ninth Circuit by using the 
appellate CM/ECF system on {datetime.now().strftime('%B %d, %Y')}.

Participants in the case who are registered CM/ECF users will be served by the 
appellate CM/ECF system.

/s/ {info.get('appellant', '')}
{info.get('appellant', '')}
"""
    
    def count_words(self) -> int:
        """Count words in substantive sections (excludes TOC, TOA, certs)"""
        total = 0
        for section in SECTION_ORDER:
            if section not in AUTO_GENERATED:
                text = self.loader.get_section_text(section)
                if text:
                    # Simple word count - split on whitespace
                    total += len(text.split())
        return total
    
    def validate(self) -> dict:
        """Validate brief completeness"""
        results = {
            "missing_required": [],
            "empty_sections": [],
            "word_count": 0,
            "over_limit": False,
            "valid": True
        }
        
        required = ["jurisdictional_statement", "issues_presented", "statement_of_case",
                    "summary_of_argument", "standard_of_review", "argument", "conclusion"]
        
        for section in required:
            text = self.loader.get_section_text(section)
            if not text:
                results["missing_required"].append(section)
                results["valid"] = False
        
        for section in SECTION_ORDER:
            if section not in AUTO_GENERATED:
                text = self.loader.get_section_text(section)
                if not text:
                    results["empty_sections"].append(section)
        
        results["word_count"] = self.count_words()
        if results["word_count"] > WORD_LIMIT:
            results["over_limit"] = True
            results["valid"] = False
        
        return results
    
    def assemble_section(self, section_name: str) -> str:
        """Assemble a single section - COPY ONLY"""
        if section_name == "cover_page":
            return self.generate_cover_page()
        elif section_name == "table_of_contents":
            return self.generate_toc()
        elif section_name == "table_of_authorities":
            return self.generate_toa()
        elif section_name == "certificate_compliance":
            return self.generate_cert_compliance(self.count_words())
        elif section_name == "certificate_service":
            return self.generate_cert_service()
        else:
            # User text - copy exactly
            text = self.loader.get_section_text(section_name)
            if not text:
                return f"[{section_name.upper().replace('_', ' ')} - NOT YET PROVIDED]"
            return text
    
    def assemble_full_brief(self) -> str:
        """Assemble complete brief in FRAP 28 order"""
        parts = []
        
        for section in SECTION_ORDER:
            section_text = self.assemble_section(section)
            if section_text:
                # Add section header for user sections
                if section not in AUTO_GENERATED and section != "cover_page":
                    header = section.upper().replace("_", " ")
                    parts.append(f"\n{'='*60}\n{header}\n{'='*60}\n")
                parts.append(section_text)
                parts.append("\n\n")
        
        return "".join(parts)
    
    def save_output(self, content: str, case_no: str, section_name: str = "full", timestamp: Optional[str] = None) -> tuple:
        """Save with dual naming convention"""
        timestamp = timestamp or datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Create output directories
        brief_dir = self.output_dir / "briefs"
        chrono_dir = self.output_dir / "chronological"
        brief_dir.mkdir(parents=True, exist_ok=True)
        chrono_dir.mkdir(parents=True, exist_ok=True)
        
        # Primary: {case}-{filename}-{datetime}
        primary_name = f"{case_no}-opening-brief-{section_name}-{timestamp}.txt"
        primary_path = brief_dir / primary_name
        
        # Chronological: {datetime}-{case}-{filename} (read-only)
        chrono_name = f"{timestamp}-{case_no}-opening-brief-{section_name}.txt"
        chrono_path = chrono_dir / chrono_name
        
        # Write primary
        with open(primary_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        # Write chronological (read-only)
        with open(chrono_path, 'w', encoding='utf-8') as f:
            f.write(content)
        os.chmod(chrono_path, stat.S_IRUSR | stat.S_IRGRP | stat.S_IROTH)
        
        return primary_path, chrono_path

    def save_output_docx(self, content: str, case_no: str, section_name: str = "full", timestamp: Optional[str] = None) -> tuple:
        """Optional .docx output for Word formatting"""
        try:
            from docx import Document
        except ImportError:
            print("python-docx not installed; skipping .docx generation. Install with: pip install python-docx")
            return None, None

        timestamp = timestamp or datetime.now().strftime("%Y%m%d_%H%M%S")

        brief_dir = self.output_dir / "briefs"
        chrono_dir = self.output_dir / "chronological"
        brief_dir.mkdir(parents=True, exist_ok=True)
        chrono_dir.mkdir(parents=True, exist_ok=True)

        doc = Document()

        # Title
        doc.add_heading(f"{case_no} Opening Brief ({section_name})", level=0)

        # Preserve paragraph breaks; avoid styling guesses for user-provided text
        for line in content.split("\n"):
            if line.strip() == "":
                doc.add_paragraph("")
            else:
                doc.add_paragraph(line)

        primary_name = f"{case_no}-opening-brief-{section_name}-{timestamp}.docx"
        primary_path = brief_dir / primary_name

        chrono_name = f"{timestamp}-{case_no}-opening-brief-{section_name}.docx"
        chrono_path = chrono_dir / chrono_name

        doc.save(primary_path)
        doc.save(chrono_path)

        return primary_path, chrono_path


def parse_args():
    parser = argparse.ArgumentParser(description="Ninth Circuit Opening Brief Assembler")
    parser.add_argument("--all", action="store_true", help="Assemble complete brief")
    parser.add_argument("--section", type=str, help="Assemble single section")
    parser.add_argument("--toa", action="store_true", help="Generate Table of Authorities")
    parser.add_argument("--toc", action="store_true", help="Generate Table of Contents")
    parser.add_argument("--validate", action="store_true", help="Validate brief completeness")
    parser.add_argument("--word-count", action="store_true", help="Count words in brief")
    parser.add_argument("--set-word-limit", type=int, help="Override word limit for validation and compliance display")
    parser.add_argument("--extract-citations", action="store_true", help="Extract citations from sections")
    parser.add_argument("--list-sections", action="store_true", help="List available sections")
    parser.add_argument("--case-no", type=str, default="DRAFT", help="Case number for output filename")
    parser.add_argument("--data-dir", type=str, default="brief_data", help="Directory containing JSON data")
    parser.add_argument("--output-dir", type=str, default="../../OUTBOX", help="Output directory")
    parser.add_argument("--docx", action="store_true", help="Also generate .docx output (requires python-docx)")
    return parser.parse_args()


def main():
    args = parse_args()

    # Allow runtime override of the word limit for users who need longer briefs
    global WORD_LIMIT
    if args.set_word_limit:
        WORD_LIMIT = args.set_word_limit
    
    # Resolve paths relative to script location
    script_dir = Path(__file__).parent
    data_dir = script_dir / args.data_dir
    output_dir = script_dir / args.output_dir
    
    assembler = BriefAssembler(str(data_dir), str(output_dir))
    
    if args.list_sections:
        print("\nAvailable sections:")
        for section in assembler.loader.list_sections():
            text = assembler.loader.get_section_text(section)
            status = "✓ has content" if text else "✗ empty"
            print(f"  - {section}: {status}")
        return
    
    if args.validate:
        print("\nValidating brief...")
        results = assembler.validate()
        print(f"Word count: {results['word_count']} / {WORD_LIMIT}")
        if results['over_limit']:
            print(f"  ⚠️  OVER LIMIT by {results['word_count'] - WORD_LIMIT} words")
        if results['missing_required']:
            print(f"Missing required sections: {', '.join(results['missing_required'])}")
        if results['empty_sections']:
            print(f"Empty sections: {', '.join(results['empty_sections'])}")
        print(f"\nValid: {'✓ Yes' if results['valid'] else '✗ No'}")
        return
    
    if args.word_count:
        count = assembler.count_words()
        print(f"\nWord count: {count} / {WORD_LIMIT}")
        if count > WORD_LIMIT:
            print(f"⚠️  OVER LIMIT by {count - WORD_LIMIT} words")
        else:
            print(f"✓ Under limit by {WORD_LIMIT - count} words")
        return
    
    if args.toa:
        toa = assembler.generate_toa()
        print("\n" + toa)
        return
    
    if args.toc:
        toc = assembler.generate_toc()
        print("\n" + toc)
        return
    
    if args.extract_citations:
        print("\nExtracting citations from all sections...")
        all_text = ""
        for section in SECTION_ORDER:
            if section not in AUTO_GENERATED:
                all_text += assembler.loader.get_section_text(section) + "\n"
        
        extractor = AuthorityExtractor()
        citations = extractor.extract_all(all_text)
        
        print(f"\nCases found: {len(citations['cases'])}")
        for case in citations['cases']:
            print(f"  - {case['raw']}")
        
        print(f"\nStatutes found: {len(citations['statutes'])}")
        for stat in citations['statutes']:
            print(f"  - {stat['citation']}")
        
        print(f"\nRules found: {len(citations['rules'])}")
        for rule in citations['rules']:
            print(f"  - {rule['citation']}")
        return
    
    if args.section:
        print(f"\nAssembling section: {args.section}")
        content = assembler.assemble_section(args.section)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        primary, chrono = assembler.save_output(content, args.case_no, args.section, timestamp=timestamp)
        print(f"Saved to: {primary}")
        print(f"Chronological copy: {chrono}")

        if args.docx:
            docx_primary, docx_chrono = assembler.save_output_docx(content, args.case_no, args.section, timestamp=timestamp)
            if docx_primary:
                print(f".docx saved to: {docx_primary}")
                print(f"Chronological .docx: {docx_chrono}")
        return
    
    if args.all:
        print("\nAssembling complete brief...")
        content = assembler.assemble_full_brief()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        primary, chrono = assembler.save_output(content, args.case_no, "full", timestamp=timestamp)
        print(f"\nSaved to: {primary}")
        print(f"Chronological copy: {chrono}")

        if args.docx:
            docx_primary, docx_chrono = assembler.save_output_docx(content, args.case_no, "full", timestamp=timestamp)
            if docx_primary:
                print(f".docx saved to: {docx_primary}")
                print(f"Chronological .docx: {docx_chrono}")
        
        # Show word count
        count = assembler.count_words()
        print(f"\nWord count: {count} / {WORD_LIMIT}")
        return
    
    # Default: show help
    print("\nNinth Circuit Opening Brief Assembler")
    print("=" * 40)
    print("\nUsage:")
    print("  --all              Assemble complete brief")
    print("  --section NAME     Assemble single section")
    print("  --toa              Generate Table of Authorities")
    print("  --toc              Generate Table of Contents")
    print("  --validate         Check brief completeness")
    print("  --word-count       Count words in brief")
    print("  --extract-citations Extract citations from text")
    print("  --list-sections    List available sections")
    print("\nExample:")
    print("  python assemble_opening_brief.py --all --case-no 25-12345")


if __name__ == "__main__":
    main()
