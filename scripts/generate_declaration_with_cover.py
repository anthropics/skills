#!/usr/bin/env python3
"""Generate a single DOCX that combines the existing cover template with a styled declaration body.

- Pulls cover styling from PIMP-SMACK-APP/_formatting/templates/TEMPLATE_CAPTION.docx
- Inserts a center header with Case No.
- Inserts a center footer with page number field
- Page 2: Heading 1 for the declaration title, preamble, numbered fact placeholders
- Facts are supplied via JSON or default placeholders; count adapts to provided facts

Run example (from repo root):
  python scripts/generate_declaration_with_cover.py \
    --case "25-6461" \
    --declarant "Tyler Allen Lofall" \
    --filing "Declaration of Tyler Allen Lofall" \
    --facts facts.json

facts.json format:
{
  "facts": [
    "I am competent to testify to these facts.",
    "I am the appellant in this matter.",
    "Primary substantive fact...",
    "Supporting detail..."
  ]
}
If --facts is omitted, five numbered placeholders are used.
"""
from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
COVER_TEMPLATE = ROOT / "PIMP-SMACK-APP" / "_formatting" / "templates" / "TEMPLATE_CAPTION.docx"
OUTBOX = ROOT / "OUTBOX"
OUTPUT_PATH = OUTBOX / "DECLARATION_WITH_COVER.docx"


def load_facts(facts_path: Path | None) -> List[str]:
    if not facts_path:
        return [
            "[Fact 1]",
            "[Fact 2]",
            "[Fact 3]",
            "[Fact 4]",
            "[Fact 5]",
        ]
    data = json.loads(facts_path.read_text(encoding="utf-8"))
    facts = data.get("facts", [])
    if not facts:
        return ["[Fact 1]"]
    return [str(f) for f in facts]


def add_header_footer(doc: Document, case_number: str):
    section = doc.sections[0]
    # Header centered
    header = section.header
    header.is_linked_to_previous = False
    if header.paragraphs:
        p = header.paragraphs[0]
        p.text = ""
    else:
        p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Case No. {case_number}")
    run.font.size = Pt(12)

    # Footer centered with page number field
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run("Page ")
    # Insert PAGE field
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r = fp._p.add_r()
    r._r.append(fld_begin)
    r._r.append(instr_text)
    r._r.append(fld_sep)
    r._r.append(fld_end)


def add_declaration_body(doc: Document, filing_title: str, declarant: str, facts: List[str]):
    doc.add_page_break()

    # Heading
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(filing_title)
    r.style = doc.styles["Heading 1"] if "Heading 1" in doc.styles else None
    r.font.bold = True
    r.font.size = Pt(14)

    # Preamble
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5)
    r = p.add_run(f"I, {declarant}, declare under penalty of perjury:")
    r.font.size = Pt(12)

    # Facts
    for idx, fact in enumerate(facts, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f"{idx}. {fact}")
        r.font.size = Pt(12)

    # Closing
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.add_run("I declare under penalty of perjury that the foregoing is true and correct.")

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.add_run("Executed on [Date], at [City, State].")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.add_run("______________________________")

    p = doc.add_paragraph()
    p.add_run(declarant)


def main():
    parser = argparse.ArgumentParser(description="Generate declaration with cover")
    parser.add_argument("--case", dest="case_number", default="[XX-XXXXX]")
    parser.add_argument("--declarant", dest="declarant", default="[Declarant Name]")
    parser.add_argument("--filing", dest="filing_title", default="Declaration")
    parser.add_argument("--facts", dest="facts_path", type=Path, default=None)
    args = parser.parse_args()

    if not COVER_TEMPLATE.exists():
        raise SystemExit(f"Cover template missing: {COVER_TEMPLATE}")

    facts = load_facts(args.facts_path)

    OUTBOX.mkdir(exist_ok=True)

    doc = Document(COVER_TEMPLATE)
    add_header_footer(doc, args.case_number)
    add_declaration_body(doc, args.filing_title, args.declarant, facts)

    doc.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
