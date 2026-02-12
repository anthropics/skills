#!/usr/bin/env python3
"""
PIMP SMACK Legal Document Formatter
====================================
Uses python-docx to create properly formatted DOCX files.
NO subprocess, NO external dependencies beyond python-docx.

Install: pip install python-docx

Usage:
  python format_document.py USER_SCHEMA.json INPUT.docx [OUTPUT.docx]
  python format_document.py --from-text INPUT.txt OUTPUT.docx
  python format_document.py --new-brief OUTPUT.docx
"""

import json
import re
import argparse
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed")
    print("Install with: pip install python-docx")
    exit(1)

# Import collector - handles case data persistence
sys.path.insert(0, str(Path(__file__).parent.parent.parent.parent))
try:
    from pimp_collector import PimpCollector
    COLLECTOR_AVAILABLE = True
except ImportError:
    COLLECTOR_AVAILABLE = False
    print("[INFO] pimp_collector not found - running without data collection")


class PimpFormatter:
    """Legal document formatter using python-docx with integrated data collection."""
    
    def __init__(self, schema_path=None, master_config_path=None):
        self.script_dir = Path(__file__).parent.parent
        self.schemas_dir = self.script_dir / "schemas"
        
        # Initialize collector for data persistence
        self.collector = None
        if COLLECTOR_AVAILABLE:
            app_dir = Path(__file__).parent.parent.parent.parent
            config_path = app_dir / "MASTER_CASE_CONFIG.json"
            self.collector = PimpCollector(str(config_path))
        
        # Load master config if provided
        self.master_config = None
        if master_config_path and Path(master_config_path).exists():
            self.master_config = self.load_json(master_config_path)
        
        # Load user schema
        if schema_path:
            self.user_schema = self.load_json(schema_path)
        else:
            self.user_schema = {}
        
        # Load and merge with master schema
        self.master_schema = self.load_master_schema()
        self.config = self.merge_schemas()
    
    def load_json(self, path):
        """Load JSON file."""
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    
    def load_master_schema(self):
        """Load the appropriate master schema from local schemas folder."""
        inherits_from = self.user_schema.get('_inherits_from', 'MASTER_FRCP')
        schema_path = self.schemas_dir / f"{inherits_from}.json"
        
        if not schema_path.exists():
            print(f"WARNING: Master schema not found: {schema_path}")
            # Try MASTER_FRCP as fallback
            schema_path = self.schemas_dir / "MASTER_FRCP.json"
            if not schema_path.exists():
                print("Using built-in defaults")
                return self.get_default_schema()
        
        return self.load_json(schema_path)
    
    def get_default_schema(self):
        """Built-in default schema if no JSON found."""
        return {
            "default_formatting": {
                "font": "Century Schoolbook",
                "font_size": "14pt",
                "line_spacing": "double"
            },
            "heading_styles": {
                "LEGAL_H1": {
                    "font": "Century Schoolbook",
                    "font_size": "14pt",
                    "bold": True,
                    "caps": True,
                    "alignment": "center",
                    "line_spacing": "single"
                },
                "LEGAL_H2": {
                    "font": "Century Schoolbook",
                    "font_size": "14pt",
                    "bold": True,
                    "alignment": "left",
                    "line_spacing": "single"
                }
            },
            "body_text_style": {
                "LEGAL_BODY": {
                    "font": "Century Schoolbook",
                    "font_size": "14pt",
                    "alignment": "left",
                    "first_line_indent": "0.5in",
                    "line_spacing": "double"
                }
            }
        }
    
    def merge_schemas(self):
        """Merge user schema over master schema."""
        config = self.master_schema.copy()
        
        # Apply formatting overrides
        if 'formatting_overrides' in self.user_schema:
            for key, value in self.user_schema['formatting_overrides'].items():
                if not key.startswith('_'):
                    if 'default_formatting' in config and key in config['default_formatting']:
                        config['default_formatting'][key] = value
        
        # Add user info
        config['filing_info'] = self.user_schema.get('filing_info', {})
        config['case_info'] = self.user_schema.get('case_info', {})
        config['headings_list'] = self.user_schema.get('headings_in_my_document', [])
        
        return config
    
    def create_legal_styles(self, doc):
        """Add LEGAL_H1, LEGAL_H2, LEGAL_BODY styles to document."""
        styles = doc.styles
        
        # Get font settings
        heading_config = self.config.get('heading_styles', {}).get('LEGAL_H1', {})
        body_config = self.config.get('body_text_style', {}).get('LEGAL_BODY', {})
        
        font_name = heading_config.get('font', 'Century Schoolbook')
        heading_size = int(heading_config.get('font_size', '14pt').replace('pt', ''))
        body_size = int(body_config.get('font_size', '14pt').replace('pt', ''))
        
        # LEGAL_H1 - Main headings (centered, bold, caps)
        try:
            h1_style = styles.add_style('LEGAL_H1', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            h1_style = styles['LEGAL_H1']
        
        h1_style.font.name = font_name
        h1_style.font.size = Pt(heading_size)
        h1_style.font.bold = True
        h1_style.font.all_caps = True
        h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h1_style.paragraph_format.space_before = Pt(0)
        h1_style.paragraph_format.space_after = Pt(0)
        h1_style.paragraph_format.line_spacing = 1.0
        
        # LEGAL_H2 - Subheadings (left, bold)
        try:
            h2_style = styles.add_style('LEGAL_H2', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            h2_style = styles['LEGAL_H2']
        
        h2_style.font.name = font_name
        h2_style.font.size = Pt(heading_size)
        h2_style.font.bold = True
        h2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h2_style.paragraph_format.space_before = Pt(0)
        h2_style.paragraph_format.space_after = Pt(0)
        h2_style.paragraph_format.line_spacing = 1.0
        
        # LEGAL_H3 - Sub-subheadings (left, bold, indented)
        try:
            h3_style = styles.add_style('LEGAL_H3', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            h3_style = styles['LEGAL_H3']
        
        h3_style.font.name = font_name
        h3_style.font.size = Pt(heading_size)
        h3_style.font.bold = True
        h3_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h3_style.paragraph_format.left_indent = Inches(0.5)
        h3_style.paragraph_format.line_spacing = 1.0
        
        # LEGAL_BODY - Body text (double-spaced, first line indent)
        try:
            body_style = styles.add_style('LEGAL_BODY', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            body_style = styles['LEGAL_BODY']
        
        body_style.font.name = font_name
        body_style.font.size = Pt(body_size)
        body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        body_style.paragraph_format.first_line_indent = Inches(0.5)
        body_style.paragraph_format.line_spacing = 2.0
        body_style.paragraph_format.space_before = Pt(0)
        body_style.paragraph_format.space_after = Pt(0)
        
        return doc
    
    def set_page_margins(self, doc):
        """Set 1-inch margins on all sides."""
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        return doc
    
    def detect_heading_level(self, text):
        """Detect if text is a heading and what level."""
        text_upper = text.strip().upper()
        
        # Check against known headings list
        headings_list = self.config.get('headings_list', [])
        for heading in headings_list:
            if text_upper == heading.upper() or text_upper == heading.replace('_', ' ').upper():
                return 'LEGAL_H1'
        
        # Check for common H1 patterns
        h1_keywords = [
            'INTRODUCTION', 'JURISDICTIONAL STATEMENT', 'STATEMENT OF ISSUES',
            'STATEMENT OF THE CASE', 'STATEMENT OF FACTS', 'SUMMARY OF ARGUMENT',
            'STANDARD OF REVIEW', 'ARGUMENT', 'CONCLUSION', 'RELATED CASES',
            'FACTUAL BACKGROUND', 'LEGAL STANDARD', 'PROCEDURAL HISTORY'
        ]
        
        if text_upper in h1_keywords:
            return 'LEGAL_H1'
        
        # Check for numbered H2 patterns (I., II., A., B.)
        import re
        if re.match(r'^[IVX]+\.\s+', text) or re.match(r'^[A-Z]\.\s+', text):
            return 'LEGAL_H2'
        
        # Check for numbered H3 patterns (1., 2., a., b.)
        if re.match(r'^\d+\.\s+', text) or re.match(r'^[a-z]\.\s+', text):
            return 'LEGAL_H3'
        
        return None
    
    def format_existing_docx(self, input_path, output_path):
        """Format an existing DOCX file - preserve text, apply styles."""
        print("\n" + "=" * 60)
        print("PIMP SMACK FORMATTER")
        print("=" * 60)
        print(f"\nInput: {input_path}")
        
        # Extract case data using collector
        if self.collector:
            print("\n[COLLECTOR] Extracting case data...")
            self.collector.extract_from_docx(input_path)
        
        doc = Document(input_path)
        doc = self.create_legal_styles(doc)
        doc = self.set_page_margins(doc)
        
        heading_count = 0
        body_count = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            heading_level = self.detect_heading_level(text)
            
            if heading_level:
                para.style = heading_level
                heading_count += 1
                print(f"  [H] {text[:50]}...")
            else:
                para.style = 'LEGAL_BODY'
                body_count += 1
        
        doc.save(output_path)
        
        # Save collected data and show status
        if self.collector:
            self.collector.save()
            stats = self.collector.get_stats()
            print(f"\n[COLLECTOR] Case: {stats['case_number']}")
            print(f"[COLLECTOR] Sections: {stats['sections_complete']} | Citations: {stats['citations_collected']}")
        
        print(f"\n{'=' * 60}")
        print(f"FORMATTED: {heading_count} headings, {body_count} body paragraphs")
        print(f"OUTPUT: {output_path}")
        print("=" * 60)
        
        return output_path
    
    def create_brief_from_text(self, input_text_path, output_path):
        """Create formatted DOCX from plain text file."""
        print("\n" + "=" * 60)
        print("PIMP SMACK - TEXT TO DOCX")
        print("=" * 60)
        
        with open(input_text_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Extract case data using collector
        if self.collector:
            print("\n[COLLECTOR] Extracting case data...")
            self.collector.extract_from_text(content)
        
        doc = Document()
        doc = self.create_legal_styles(doc)
        doc = self.set_page_margins(doc)
        
        heading_count = 0
        body_count = 0
        
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                doc.add_paragraph('')
                continue
            
            heading_level = self.detect_heading_level(line)
            
            if heading_level:
                para = doc.add_paragraph(line, style=heading_level)
                heading_count += 1
                print(f"  [H] {line[:50]}...")
            else:
                para = doc.add_paragraph(line, style='LEGAL_BODY')
                body_count += 1
        
        doc.save(output_path)
        
        # Save collected data and show status
        if self.collector:
            self.collector.save()
            stats = self.collector.get_stats()
            print(f"\n[COLLECTOR] Case: {stats['case_number']}")
            print(f"[COLLECTOR] Sections: {stats['sections_complete']} | Citations: {stats['citations_collected']}")
        
        print(f"\n{'=' * 60}")
        print(f"CREATED: {heading_count} headings, {body_count} body paragraphs")
        print(f"OUTPUT: {output_path}")
        print("=" * 60)
        
        return output_path
    
    def create_new_brief(self, output_path, sections=None):
        """Create a new formatted brief from section data."""
        print("\n" + "=" * 60)
        print("PIMP SMACK - NEW BRIEF GENERATOR")
        print("=" * 60)
        
        doc = Document()
        doc = self.create_legal_styles(doc)
        doc = self.set_page_margins(doc)
        
        # Default sections if none provided
        if sections is None:
            sections = {
                "INTRODUCTION": "[Your introduction here]",
                "JURISDICTIONAL STATEMENT": "[Your jurisdictional statement here]",
                "STATEMENT OF ISSUES": "[Your issues presented here]",
                "STATEMENT OF THE CASE": "[Your statement of case here]",
                "SUMMARY OF ARGUMENT": "[Your summary here]",
                "STANDARD OF REVIEW": "[Your standard of review here]",
                "ARGUMENT": "[Your argument here]",
                "CONCLUSION": "[Your conclusion here]"
            }
        
        for heading, content in sections.items():
            # Add heading
            doc.add_paragraph(heading, style='LEGAL_H1')
            
            # Add content
            for para in content.split('\n\n'):
                if para.strip():
                    doc.add_paragraph(para.strip(), style='LEGAL_BODY')
        
        doc.save(output_path)
        
        print(f"\nCREATED: {output_path}")
        print("=" * 60)
        
        return output_path


def main():
    parser = argparse.ArgumentParser(
        description='PIMP SMACK Legal Document Formatter',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  Format existing DOCX:
    python format_document.py schema.json input.docx output.docx
  
  Convert text to DOCX:
    python format_document.py --from-text input.txt output.docx
  
  Create new brief template:
    python format_document.py --new-brief output.docx
        """
    )
    
    parser.add_argument('schema', nargs='?', help='User schema JSON file')
    parser.add_argument('input', nargs='?', help='Input DOCX or TXT file')
    parser.add_argument('output', nargs='?', help='Output DOCX file')
    parser.add_argument('--from-text', action='store_true', help='Convert text file to DOCX')
    parser.add_argument('--new-brief', action='store_true', help='Create new brief template')
    parser.add_argument('--master-config', type=str, help='Path to MASTER_CASE_CONFIG.json')
    
    args = parser.parse_args()
    
    # Handle --new-brief
    if args.new_brief:
        output = args.schema or f"new_brief_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        formatter = PimpFormatter()
        formatter.create_new_brief(output)
        return
    
    # Handle --from-text
    if args.from_text:
        if not args.schema or not args.input:
            print("Usage: python format_document.py --from-text INPUT.txt OUTPUT.docx")
            exit(1)
        formatter = PimpFormatter()
        formatter.create_brief_from_text(args.schema, args.input)
        return
    
    # Standard format existing DOCX
    if not args.schema or not args.input:
        parser.print_help()
        exit(1)
    
    output = args.output or f"{Path(args.input).stem}_FORMATTED_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    formatter = PimpFormatter(args.schema, args.master_config)
    formatter.format_existing_docx(args.input, output)


if __name__ == '__main__':
    main()
