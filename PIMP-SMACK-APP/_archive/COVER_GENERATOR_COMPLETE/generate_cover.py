#!/usr/bin/env python3
"""
Ninth Circuit Cover Page Generator
Keeps the master template pristine and generates new covers by swapping placeholders

USAGE:
  Interactive:  python generate_cover.py
  CLI (Claude): python generate_cover.py --case "25-6461" --filing "APPELLANT'S OPENING BRIEF" --judge "Stacy Beckerman"
  
  Optional: --output "custom_filename.docx"
"""

import zipfile
import os
import sys
import shutil
import argparse
from datetime import datetime
from pathlib import Path


def get_unique_filename(base_path: str) -> str:
    """
    Check if file exists, add number suffix if needed.
    Example: file.docx -> file_1.docx -> file_2.docx
    """
    path = Path(base_path)
    if not path.exists():
        return str(path)
    
    stem = path.stem
    suffix = path.suffix
    parent = path.parent
    
    counter = 1
    while True:
        new_path = parent / f"{stem}_{counter}{suffix}"
        if not new_path.exists():
            return str(new_path)
        counter += 1


def parse_args():
    """Parse command-line arguments for non-interactive mode"""
    parser = argparse.ArgumentParser(
        description='Generate Ninth Circuit Cover Page',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python generate_cover.py --case "25-6461" --filing "APPELLANT'S OPENING BRIEF" --judge "Stacy Beckerman"
  python generate_cover.py --case "25-6461" --filing "MOTION FOR STAY" --judge "Ann Aiken" --output "my_cover.docx"
        """
    )
    parser.add_argument('--case', type=str, help='Ninth Circuit case number (e.g., 25-6461)')
    parser.add_argument('--filing', type=str, help='Filing name (e.g., APPELLANT\'S OPENING BRIEF)')
    parser.add_argument('--judge', type=str, help='District judge name (e.g., Stacy Beckerman)')
    parser.add_argument('--output', type=str, help='Output filename (optional, auto-generated if not provided)')
    parser.add_argument('--template', type=str, help='Path to template file (optional)')
    parser.add_argument('--font', type=str, help='Font name to apply (e.g., Century Schoolbook)')
    parser.add_argument('--size', type=float, help='Font size to apply (e.g., 13)')
    
    return parser.parse_args()


def prompt_for_values():
    """Prompt user for all placeholder values (interactive mode)"""
    print("\n" + "="*60)
    print("NINTH CIRCUIT COVER PAGE GENERATOR")
    print("="*60 + "\n")
    
    # Case number (Ninth Circuit)
    print("Enter Ninth Circuit case number (or press Enter for blank):")
    print("  Example: 24-1234")
    case_number = input("  Case #: ").strip()
    if not case_number:
        case_number = "____________________"
    else:
        case_number = f"No. {case_number}"
    
    # Filing name
    print("\nEnter filing name:")
    print("  Examples:")
    print("    APPELLANT'S OPENING BRIEF")
    print("    APPELLANT'S REPLY BRIEF")
    print("    MOTION FOR STAY PENDING APPEAL")
    filing_name = input("  Filing: ").strip().upper()
    if not filing_name:
        filing_name = "APPELLANT'S OPENING BRIEF"
    
    # Judge name
    print("\nEnter district judge name (or press Enter for placeholder):")
    print("  Example: Stacy Beckerman")
    judge_name = input("  Judge: ").strip()
    if not judge_name:
        judge_name = "[District Judge Name]"
    else:
        judge_name = f"Hon. {judge_name}"
    
    print("\n" + "="*60)
    print("GENERATING COVER PAGE...")
    print("="*60 + "\n")
    
    return {
        'case_number': case_number,
        'filing_name': filing_name,
        'judge_name': judge_name
    }


def values_from_args(args):
    """Build values dict from command-line arguments (non-interactive mode)"""
    case_number = args.case if args.case else "____________________"
    if args.case and not args.case.startswith("No."):
        case_number = f"No. {args.case}"
    
    filing_name = args.filing.upper() if args.filing else "APPELLANT'S OPENING BRIEF"
    
    judge_name = args.judge if args.judge else "[District Judge Name]"
    if args.judge and not args.judge.startswith("Hon."):
        judge_name = f"Hon. {args.judge}"
    
    return {
        'case_number': case_number,
        'filing_name': filing_name,
        'judge_name': judge_name,
        'font_name': args.font,
        'font_size': args.size
    }

def generate_cover(template_path, output_path, values):
    """
    Generate a new cover page from the template by replacing placeholders
    
    Args:
        template_path: Path to the master template (TEMPLATE_CAPTION.docx)
        output_path: Path for the generated file
        values: Dictionary with placeholder values
    """
    
    # Create a temporary directory for extraction (Windows compatible)
    import tempfile
    temp_dir = os.path.join(tempfile.gettempdir(), "cover_temp")
    if os.path.exists(temp_dir):
        shutil.rmtree(temp_dir)
    os.makedirs(temp_dir)
    
    # Extract the template docx (it's a ZIP file)
    with zipfile.ZipFile(template_path, 'r') as zip_ref:
        zip_ref.extractall(temp_dir)
    
    # Read the document.xml
    doc_xml_path = os.path.join(temp_dir, 'word', 'document.xml')
    with open(doc_xml_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Replace placeholders
    # Case number
    content = content.replace('No. 6461', values['case_number'])
    
    # Filing name (in FILLIN field)
    content = content.replace('APPELLANTS OPENING BRIEF', values['filing_name'])
    
    # Judge name
    content = content.replace('Hon. Stacy Beckerman', values['judge_name'])
    
    # Write back the modified XML
    with open(doc_xml_path, 'w', encoding='utf-8') as f:
        f.write(content)
    
    # Re-package as a .docx file
    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as docx:
        for foldername, subfolders, filenames in os.walk(temp_dir):
            for filename in filenames:
                file_path = os.path.join(foldername, filename)
                arcname = os.path.relpath(file_path, temp_dir)
                docx.write(file_path, arcname)
    
    # Clean up temp directory
    shutil.rmtree(temp_dir)

    # Apply formatting if requested
    if values.get('font_name') or values.get('font_size'):
        try:
            from docx import Document
            from docx.shared import Pt
            
            doc = Document(output_path)
            font_name = values.get('font_name')
            font_size = values.get('font_size')
            
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if font_name:
                        run.font.name = font_name
                    if font_size:
                        run.font.size = Pt(float(font_size))
            
            # Also check tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if font_name:
                                    run.font.name = font_name
                                if font_size:
                                    run.font.size = Pt(float(font_size))
            
            doc.save(output_path)
            print(f"[OK] Applied formatting: {font_name} {font_size}pt")
        except ImportError:
            print("[WARN] python-docx not installed, skipping formatting application")
        except Exception as e:
            print(f"[WARN] Failed to apply formatting: {e}")
    
    print(f"[OK] Cover page generated: {output_path}")
    print(f"  Case Number: {values['case_number']}")
    print(f"  Filing Name: {values['filing_name']}")
    print(f"  Judge: {values['judge_name']}")


def make_readonly(filepath: str):
    """Make a file read-only"""
    import stat
    os.chmod(filepath, stat.S_IREAD)


def save_dual_copies(source_path: str, case_num: str, filename: str, outbox_dir: Path, section_subdir: str):
    """
    Save two copies:
    1. Primary:      {case_no}-{filename}-{datetime}.docx  (in section subdir)
    2. Chronological: {datetime}-{case_no}-{filename}.docx (read-only, in chronological subdir)
    
    Returns tuple of (primary_path, chrono_path)
    """
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Clean names for filesystem
    clean_case = case_num.replace("No. ", "").replace(" ", "").replace("/", "-")
    clean_filename = filename.replace("'", "").replace(" ", "_").replace("/", "-")
    
    # Create section subdirectory
    section_dir = outbox_dir / section_subdir
    section_dir.mkdir(parents=True, exist_ok=True)
    
    # Create chronological subdirectory
    chrono_dir = outbox_dir / "chronological"
    chrono_dir.mkdir(parents=True, exist_ok=True)
    
    # Primary naming: {case_no}-{filename}-{datetime}.docx
    primary_name = f"{clean_case}-{clean_filename}-{timestamp}.docx"
    primary_path = get_unique_filename(str(section_dir / primary_name))
    
    # Chronological naming: {datetime}-{case_no}-{filename}.docx (read-only)
    chrono_name = f"{timestamp}-{clean_case}-{clean_filename}.docx"
    chrono_path = get_unique_filename(str(chrono_dir / chrono_name))
    
    # Copy to primary location
    shutil.copy2(source_path, primary_path)
    
    # Copy to chronological location and make read-only
    shutil.copy2(source_path, chrono_path)
    make_readonly(chrono_path)
    
    # Remove the temp source
    os.remove(source_path)
    
    return primary_path, chrono_path


def main():
    """Main function - supports both interactive and CLI modes"""
    
    args = parse_args()
    
    # Determine if running in CLI mode (any argument provided)
    cli_mode = args.case or args.filing or args.judge or args.output
    
    # Get script directory for template path
    script_dir = Path(__file__).parent
    
    # OUTBOX: Single location for all outputs (outside codebase)
    outbox_dir = script_dir.parent / "OUTBOX"
    outbox_dir.mkdir(exist_ok=True)
    
    # Path to the master template (READ-ONLY)
    if args.template:
        template_path = Path(args.template)
    else:
        template_path = script_dir / "TEMPLATE_CAPTION.docx"
    
    # Check if template exists
    if not template_path.exists():
        print(f"ERROR: Template not found: {template_path}")
        print("Please ensure TEMPLATE_CAPTION.docx is in the script directory.")
        sys.exit(1)
    
    # Get values - CLI or interactive
    if cli_mode:
        values = values_from_args(args)
    else:
        values = prompt_for_values()
    
    # Get clean case number for naming
    clean_case = values['case_number'].replace("No. ", "").replace(" ", "")
    
    # Generate to temp file first
    import tempfile
    temp_output = os.path.join(tempfile.gettempdir(), "cover_temp_output.docx")
    
    # Generate the cover page to temp location
    generate_cover(str(template_path), temp_output, values)
    
    # Save dual copies with proper naming
    primary_path, chrono_path = save_dual_copies(
        source_path=temp_output,
        case_num=clean_case,
        filename=values['filing_name'] + "_COVER",
        outbox_dir=outbox_dir,
        section_subdir="covers"
    )
    
    if not cli_mode:
        print(f"\n{'='*60}")
        print("DONE! Your cover page is ready.")
        print(f"{'='*60}\n")
    
    print(f"\nOutput files:")
    print(f"  Primary:       {primary_path}")
    print(f"  Chronological: {chrono_path} (read-only)")
    
    if cli_mode and args.output:
        output_path = Path(args.output)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(primary_path, output_path)
        print(f"  Custom Output: {output_path}")

    if not cli_mode:
        print("\nNext steps:")
        print("  1. Open the file to verify it looks correct")
        print("  2. Export to PDF")
        print("  3. Combine with your body text PDF")
        print("  4. File with Ninth Circuit\n")
    
    return primary_path

if __name__ == "__main__":
    main()
