#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
月次レポート自動生成スクリプト（Jony Iveデザイン - バランス版）
2カラムレイアウトで視覚的にバランスの取れたA4縦1ページレポート
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO

# 日本語フォント設定
plt.rcParams['font.sans-serif'] = ['Noto Sans CJK JP', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

# レポート設定
REPORT_CONFIG = {
    "client_name": "合同会社大吉",
    "consultant_name": "株式会社けんけん",
    "report_month": datetime.now().strftime("%Y年%m月"),
}

def set_font(run, font_size=10, color_rgb=(51, 51, 51), bold=False):
    """フォントを設定"""
    run.font.name = 'メイリオ'
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(*color_rgb)
    run.font.bold = bold

def set_cell_border(cell):
    """セルの罫線を削除（見えないテーブル）"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_element = OxmlElement(f'w:{edge}')
        edge_element.set(qn('w:val'), 'nil')
        tcBorders.append(edge_element)
    tcPr.append(tcBorders)

def add_thin_line(paragraph):
    """薄いグレーの下線を追加"""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E0E0E0')
    pBdr.append(bottom)
    pPr.append(pBdr)

def load_daily_reports_from_gdrive(month_str):
    """Google Driveから日報データを取得（TODO）"""
    data = {
        "日付": [f"2025-01-{i:02d}" for i in range(1, 21)],
        "作業内容": [
            "戦略策定ミーティング", "AI研修準備", "月次レポート作成",
            "定例会議", "AI研修実施", "資料作成"
        ] * 3 + ["相談対応", "分析業務"],
    }
    return pd.DataFrame(data)

def analyze_daily_reports(df):
    """日報データを分析"""
    working_days = len(df)
    
    task_categories = {
        "戦略策定": 0,
        "研修": 0,
        "資料作成": 0,
        "定例会議": 0,
        "その他": 0
    }
    
    for content in df["作業内容"]:
        if "戦略" in content or "ミーティング" in content:
            task_categories["戦略策定"] += 1
        elif "研修" in content:
            task_categories["研修"] += 1
        elif "レポート" in content or "資料" in content:
            task_categories["資料作成"] += 1
        elif "会議" in content or "定例" in content:
            task_categories["定例会議"] += 1
        else:
            task_categories["その他"] += 1
    
    top_tasks = df["作業内容"].value_counts().head(3).to_dict()
    
    return {
        "working_days": working_days,
        "task_categories": task_categories,
        "top_tasks": top_tasks
    }

def create_compact_pie_chart(task_categories):
    """小さな円グラフを作成（カラー版）"""
    fig, ax = plt.subplots(figsize=(3, 3), facecolor='white')
    
    labels = []
    sizes = []
    for task, count in task_categories.items():
        if count > 0:
            labels.append(task)
            sizes.append(count)
    
    # ビジネス向けカラーパレット（落ち着いた色調）
    colors = ['#333333', '#666666', '#999999', '#CCCCCC', '#5B7B94', '#E0E0E0']
    
    ax.pie(sizes, labels=labels, autopct='%1.0f%%', startangle=90, 
           colors=colors, textprops={'fontsize': 9})
    ax.axis('equal')
    
    img_stream = BytesIO()
    plt.savefig(img_stream, format='png', dpi=150, bbox_inches='tight', 
                facecolor='white', pad_inches=0.05)
    img_stream.seek(0)
    plt.close()
    
    return img_stream

def create_monthly_report(client_name=None, report_month=None):
    """月次レポートを作成（2ページきっちり版）"""
    if client_name is None:
        client_name = REPORT_CONFIG["client_name"]
    if report_month is None:
        report_month = REPORT_CONFIG["report_month"]
    
    # データ取得・分析
    df = load_daily_reports_from_gdrive(report_month)
    analysis = analyze_daily_reports(df)
    
    # 稼働時間計算（1日8時間想定）
    working_hours = analysis['working_days'] * 8
    
    # Wordドキュメント作成
    doc = Document()
    
    # ページ設定
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ========================================
    # 1ページ目：概要
    # ========================================
    
    # === ヘッダー ===
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p.add_run("月次活動レポート")
    set_font(run1, font_size=22, color_rgb=(0, 0, 0), bold=True)
    p.paragraph_format.space_after = Pt(2)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{report_month}  |  {client_name} 様")
    set_font(run, font_size=11, color_rgb=(102, 102, 102))
    p.paragraph_format.space_after = Pt(8)
    
    # セパレーター
    p_sep = doc.add_paragraph()
    add_thin_line(p_sep)
    p_sep.paragraph_format.space_after = Pt(14)
    
    # === 稼働実績 ===
    p = doc.add_paragraph()
    run = p.add_run("稼働実績")
    set_font(run, font_size=13, color_rgb=(102, 102, 102), bold=True)
    p.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    run = p.add_run(f"稼働日数：{analysis['working_days']}日  |  稼働時間：{working_hours}時間（月額契約範囲内）")
    set_font(run, font_size=9.5, color_rgb=(51, 51, 51))
    p.paragraph_format.space_after = Pt(14)
    
    # === 活動内訳（2カラム） ===
    p = doc.add_paragraph()
    run = p.add_run("活動内訳")
    set_font(run, font_size=13, color_rgb=(102, 102, 102), bold=True)
    p.paragraph_format.space_after = Pt(8)
    
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    
    table.columns[0].width = Cm(9)
    table.columns[1].width = Cm(7)
    
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
    
    # 左：業務リスト
    left_cell = table.rows[0].cells[0]
    
    total = sum(analysis['task_categories'].values())
    for task, count in analysis['task_categories'].items():
        if count > 0:
            percentage = int((count / total) * 100)
            p = left_cell.add_paragraph()
            run = p.add_run(f"• {task}：{percentage}%（{count}日）")
            set_font(run, font_size=9.5, color_rgb=(51, 51, 51))
            p.paragraph_format.space_after = Pt(3)
    
    # 右：グラフ
    right_cell = table.rows[0].cells[1]
    img_stream = create_compact_pie_chart(analysis['task_categories'])
    p_img = right_cell.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(img_stream, width=Cm(6.5))
    
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_after = Pt(10)
    
    # === 主要活動 ===
    p = doc.add_paragraph()
    run = p.add_run("主要活動")
    set_font(run, font_size=13, color_rgb=(102, 102, 102), bold=True)
    p.paragraph_format.space_after = Pt(8)
    
    # 活動1
    p = doc.add_paragraph()
    run = p.add_run("1. 生成AI活用戦略の策定")
    set_font(run, font_size=10, color_rgb=(51, 51, 51), bold=True)
    p.paragraph_format.space_after = Pt(3)
    
    activities_1 = [
        "現状分析と課題の整理",
        "ユースケースの検討と優先順位付け",
        "6ヶ月のロードマップ作成"
    ]
    for activity in activities_1:
        p = doc.add_paragraph()
        run = p.add_run(f"   - {activity}")
        set_font(run, font_size=9, color_rgb=(51, 51, 51))
        p.paragraph_format.space_after = Pt(1)
    
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_after = Pt(4)
    
    # 活動2
    p = doc.add_paragraph()
    run = p.add_run("2. 従業員向けAI研修の実施")
    set_font(run, font_size=10, color_rgb=(51, 51, 51), bold=True)
    p.paragraph_format.space_after = Pt(3)
    
    activities_2 = [
        "基礎研修の実施（参加者20名）",
        "実践ワークショップの開催",
        "個別相談対応"
    ]
    for activity in activities_2:
        p = doc.add_paragraph()
        run = p.add_run(f"   - {activity}")
        set_font(run, font_size=9, color_rgb=(51, 51, 51))
        p.paragraph_format.space_after = Pt(1)
    
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_after = Pt(4)
    
    # 活動3
    p = doc.add_paragraph()
    run = p.add_run("3. 月次レポート・資料作成")
    set_font(run, font_size=10, color_rgb=(51, 51, 51), bold=True)
    p.paragraph_format.space_after = Pt(3)
    
    activities_3 = [
        "月次活動レポートの作成",
        "社内向け資料の作成",
        "進捗報告資料の準備"
    ]
    for activity in activities_3:
        p = doc.add_paragraph()
        run = p.add_run(f"   - {activity}")
        set_font(run, font_size=9, color_rgb=(51, 51, 51))
        p.paragraph_format.space_after = Pt(1)
    
    # ========================================
    # 2ページ目：成果・課題・次月予定
    # ========================================
    doc.add_page_break()
    
    # === 成果 ===
    p = doc.add_paragraph()
    run = p.add_run("成果")
    set_font(run, font_size=13, color_rgb=(102, 102, 102), bold=True)
    p.paragraph_format.space_after = Pt(8)
    
    achievements = [
        ("AI活用戦略の策定完了", [
            "3つの重点ユースケースを特定",
            "6ヶ月のロードマップを作成",
            "経営陣への報告・承認完了"
        ]),
        ("従業員研修の実施", [
            "参加率95%（20/21名）を達成",
            "研修満足度4.5/5.0を獲得",
            "実践課題の提出率80%"
        ]),
        ("業務自動化プランの提示", [
            "月間40時間の削減見込み",
            "3業務の自動化を提案",
            "ROI試算を提示"
        ])
    ]
    
    for title, details in achievements:
        p = doc.add_paragraph()
        run = p.add_run(f"✓ {title}")
        set_font(run, font_size=10, color_rgb=(51, 51, 51), bold=True)
        p.paragraph_format.space_after = Pt(3)
        
        for detail in details:
            p = doc.add_paragraph()
            run = p.add_run(f"   → {detail}")
            set_font(run, font_size=9, color_rgb=(51, 51, 51))
            p.paragraph_format.space_after = Pt(1)
        
        p_sep = doc.add_paragraph()
        p_sep.paragraph_format.space_after = Pt(4)
    
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_after = Pt(6)
    
    # === 課題と対応 ===
    p = doc.add_paragraph()
    run = p.add_run("課題と対応")
    set_font(run, font_size=13, color_rgb=(102, 102, 102), bold=True)
    p.paragraph_format.space_after = Pt(8)
    
    issues = [
        ("一部部署での活用が遅延", "追加説明会を次月第1週に実施予定"),
        ("セキュリティポリシーが未整備", "次月中にガイドライン策定を支援"),
        ("効果測定指標の設定が必要", "KPI設定のワークショップを実施予定")
    ]
    
    for issue, action in issues:
        p = doc.add_paragraph()
        run = p.add_run(f"• {issue}")
        set_font(run, font_size=9.5, color_rgb=(51, 51, 51))
        p.paragraph_format.space_after = Pt(2)
        
        p = doc.add_paragraph()
        run = p.add_run(f"   → {action}")
        set_font(run, font_size=9, color_rgb=(102, 102, 102))
        p.paragraph_format.space_after = Pt(5)
    
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_after = Pt(6)
    
    # === 次月の計画 ===
    p = doc.add_paragraph()
    run = p.add_run("次月の計画")
    set_font(run, font_size=13, color_rgb=(102, 102, 102), bold=True)
    p.paragraph_format.space_after = Pt(8)
    
    plans = [
        "AI活用の効果測定・分析（KPI設定と測定開始）",
        "フォローアップ研修の実施（希望者向け）",
        "セキュリティガイドラインの策定支援",
        "追加ユースケースの検討",
        "月次定例会議での進捗報告"
    ]
    
    for i, plan in enumerate(plans, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {plan}")
        set_font(run, font_size=9.5, color_rgb=(51, 51, 51))
        p.paragraph_format.space_after = Pt(3)
    
    # === フッター（余白なし） ===
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_after = Pt(10)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(REPORT_CONFIG["consultant_name"])
    set_font(run, font_size=8, color_rgb=(153, 153, 153))
    
    # ファイル保存
    filename = f"/mnt/user-data/outputs/月次レポート_{report_month.replace('年', '').replace('月', '')}.docx"
    doc.save(filename)
    print(f"月次レポートを作成しました: {filename}")
    
    return filename

if __name__ == "__main__":
    create_monthly_report()
